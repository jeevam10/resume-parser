import os
import re
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from collections import defaultdict

import pdfplumber
import fitz  # PyMuPDF
import PyPDF2
from docx import Document
import mammoth
import langchain
from langchain.text_splitter import CharacterTextSplitter

# OpenAI integration
import openai
from openai import OpenAI

from sentence_transformers import SentenceTransformer, util

from pymongo import MongoClient

# Use chunk_size suitable for a longer LLM token limit (OpenAI models ~4k tokens)
text_splitter = CharacterTextSplitter(chunk_size=900, chunk_overlap=120)
MONGO_URI = "mongodb+srv://philosopher211004:KlphBlH6x1Ixdjj1@cluster0.uvkdgmg.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0"

# Initialize OpenAI client
OPENAI_API_KEY = "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"
openai_client = OpenAI(api_key=OPENAI_API_KEY)

client = MongoClient(MONGO_URI)
db = client["resume_parser_db"]  # You can name this anything
collection = db["parsed_resumes"]
graphs_collection = db["graphs_llm"]


# Collection to store parsed resume results

class DocumentToJSONConverter:
    def __init__(self):
        self.supported_formats = ['.pdf', '.doc', '.docx']
        self.extraction_methods = {
            'pdf': ['pdfplumber', 'pymupdf', 'pypdf2'],
            'doc': ['python-docx', 'mammoth']
        }

    def convert_to_json(self, file_path: str, output_path: Optional[str] = None) -> Dict:
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")

            file_extension = Path(file_path).suffix.lower()
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")

            if file_extension == '.pdf':
                result = self._process_pdf(file_path)
            elif file_extension in ['.doc', '.docx']:
                result = self._process_doc(file_path)

            result['metadata'] = self._get_file_metadata(file_path)

            if output_path:
                self._save_json(result, output_path)

            return result

        except Exception as e:
            return self._create_error_response(str(e), file_path)

    def _process_pdf(self, file_path: str) -> Dict:
        """
        Runs all three extractors for PDFs,
        merges their raw_text outputs,
        deduplicates repeated lines,
        and attempts to group content by common resume sections.
        """
        extractor_results = []
        for method in self.extraction_methods['pdf']:
            if method == 'pdfplumber':
                content = self._extract_with_pdfplumber(file_path)
            elif method == 'pymupdf':
                content = self._extract_with_pymupdf(file_path)
            elif method == 'pypdf2':
                content = self._extract_with_pypdf2(file_path)
            if content and content.get('raw_text', '').strip():
                extractor_results.append((method, content))

        if not extractor_results:
            return {
                "content_type": "pdf",
                "extraction_method": None,
                "pages": [],
                "raw_text": "",
                "structured_content": {},
                "extraction_success": False,
                "errors": ["No extraction method succeeded"]
            }

        # Choose extractor with longest raw_text as base
        best_method, best_content = max(
            extractor_results,
            key=lambda x: len(x[1]['raw_text'].strip())
        )

        # Merge and deduplicate lines from all extractors
        all_texts = [r[1]['raw_text'] for r in extractor_results]
        merged_text = self._merge_and_deduplicate_texts(all_texts)

        # Group into logical sections based on common headers
        cleaned_text = self._section_cleanup(merged_text)

        result = dict(best_content)
        result['raw_text'] = cleaned_text
        result['extraction_method'] = best_method
        result['extraction_success'] = True
        return result

    def _merge_and_deduplicate_texts(self, text_list):
        """
        Flatten multiple texts into ordered lines,
        remove duplicate lines preserving original order.
        """
        seen = set()
        unique_lines = []
        for text in text_list:
            for line in text.splitlines():
                line_stripped = line.strip()
                if line_stripped and line_stripped not in seen:
                    seen.add(line_stripped)
                    unique_lines.append(line)
        return "\n".join(unique_lines)

    def _section_cleanup(self, text):
        """
        Group lines by common resume sections, reorder sections logically,
        and add clear section headers.
        """
        section_headers = [
            "contact", "summary", "education", "experience",
            "projects", "skills", "certifications", "languages", "technologies", "other"
        ]
        lines = text.splitlines()
        sections = {}
        current_section = "other"
        sections[current_section] = []

        for line in lines:
            lowered = line.lower().strip()
            matched_section = None
            for header in section_headers:
                if re.search(rf'\b{header}\b', lowered):
                    matched_section = header
                    break
            if matched_section:
                current_section = matched_section
                if current_section not in sections:
                    sections[current_section] = []
            sections[current_section].append(line)

        # Desired section order for output
        order = [
            "contact", "summary", "education", "experience",
            "projects", "skills", "languages", "technologies",
            "certifications", "other"
        ]
        output_lines = []
        for sect in order:
            if sect in sections and sections[sect]:
                output_lines.append(f"\n=== {sect.upper()} ===")
                output_lines.extend(sections[sect])

        return "\n".join(output_lines).strip()

    def merge_and_deduplicate_texts(self, text_list):
        """
        Merges multiple raw_texts, preserves order, deduplicates lines/blocks.
        """
        # Split by lines, preserve all
        lines = [line for text in text_list for line in text.splitlines()]
        seen = set()
        unique_lines = []
        for line in lines:
            line_ = line.strip()
            if line_ and line_ not in seen:
                seen.add(line_)
                unique_lines.append(line_)
        return "\n".join(unique_lines)

    def _section_cleanup(self, text):
        """
        Detects and groups classical resume sections.
        """
        headers = ["education", "experience", "projects", "skills", "certifications", "contact", "summary"]
        lines = text.splitlines()
        sections = {}
        current_header = "other"
        sections[current_header] = []
        for line in lines:
            for h in headers:
                if h in line.lower():
                    current_header = h
                    sections.setdefault(current_header, [])
                    break
            sections[current_header].append(line)
        # Order sections by header importance:
        ordered_sections = ['contact', 'summary', 'education', 'experience', 'projects', 'skills', 'certifications',
                            'other']
        rebuilt = []
        for h in ordered_sections:
            if h in sections and sections[h]:
                rebuilt += [f"\n=== {h.upper()} ==="] + sections[h]
        return "\n".join(rebuilt)

    def _extract_with_pdfplumber(self, file_path: str) -> Dict:
        content = {
            "pages": [],
            "raw_text": "",
            "structured_content": {"tables": [], "text_blocks": []}
        }

        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_content = {
                    "page_number": page_num,
                    "text": page.extract_text() or "",
                    "tables": [],
                    "bbox": page.bbox
                }

                tables = page.extract_tables()
                if tables:
                    for table_idx, table in enumerate(tables):
                        page_content["tables"].append({
                            "table_id": f"page_{page_num}_table_{table_idx}",
                            "data": table,
                            "rows": len(table),
                            "columns": len(table[0]) if table else 0
                        })

                content["pages"].append(page_content)
                content["raw_text"] += page_content["text"] + "\n"

        return content

    def _extract_with_pymupdf(self, file_path: str) -> Dict:
        content = {
            "pages": [],
            "raw_text": "",
            "structured_content": {"images": [], "text_blocks": []}
        }

        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            page_content = {
                "page_number": page_num + 1,
                "text": page_text,
                "images": [],
                "text_blocks": []
            }

            blocks = page.get_text("dict")
            for block in blocks.get("blocks", []):
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        page_content["text_blocks"].append({
                            "text": span.get("text", ""),
                            "font": span.get("font", ""),
                            "size": span.get("size", 0),
                            "bbox": span.get("bbox", [])
                        })

            content["pages"].append(page_content)
            content["raw_text"] += page_text + "\n"

        doc.close()
        return content

    def _extract_with_pypdf2(self, file_path: str) -> Dict:
        content = {"pages": [], "raw_text": "", "structured_content": {}}

        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                content["pages"].append({
                    "page_number": page_num,
                    "text": page_text
                })
                content["raw_text"] += page_text + "\n"

        return content

    def _process_doc(self, file_path: str) -> Dict:
        """
        Runs all extractors for DOC/DOCX,
        merges their raw_text outputs,
        deduplicates repeated lines,
        and attempts to group content by common resume sections.
        """
        extractor_results = []
        for method in self.extraction_methods['doc']:
            if method == 'python-docx':
                content = self._extract_with_python_docx(file_path)
            elif method == 'mammoth':
                content = self._extract_with_mammoth(file_path)
            if content and content.get('raw_text', '').strip():
                extractor_results.append((method, content))

        if not extractor_results:
            return {
                "content_type": "doc",
                "extraction_method": None,
                "raw_text": "",
                "structured_content": {},
                "extraction_success": False,
                "errors": ["No extraction method succeeded"]
            }

        # Choose extractor with longest raw_text as base
        best_method, best_content = max(
            extractor_results,
            key=lambda x: len(x[1]['raw_text'].strip())
        )

        # Merge and deduplicate lines from all extractors
        all_texts = [r[1]['raw_text'] for r in extractor_results]
        merged_text = self._merge_and_deduplicate_texts(all_texts)

        # Group into logical sections based on common headers
        cleaned_text = self._section_cleanup(merged_text)

        result = dict(best_content)
        result['raw_text'] = cleaned_text
        result['extraction_method'] = best_method
        result['extraction_success'] = True
        return result

    def _extract_with_python_docx(self, file_path: str) -> Dict:
        content = {
            "raw_text": "",
            "structured_content": {"paragraphs": [], "tables": []}
        }

        doc = Document(file_path)
        for para in doc.paragraphs:
            content["structured_content"]["paragraphs"].append(para.text)
            content["raw_text"] += para.text + "\n"

        for table in doc.tables:
            table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            content["structured_content"]["tables"].append(table_data)

        return content

    def _extract_with_mammoth(self, file_path: str) -> Dict:
        content = {
            "raw_text": "",
            "structured_content": {"html": ""}
        }

        with open(file_path, "rb") as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            content["raw_text"] = result.value

        with open(file_path, "rb") as docx_file:
            html_result = mammoth.convert_to_html(docx_file)
            content["structured_content"]["html"] = html_result.value

        return content

    def _get_file_metadata(self, file_path: str) -> Dict:
        file_stat = os.stat(file_path)
        return {
            "file_name": os.path.basename(file_path),
            "file_path": file_path,
            "file_size": file_stat.st_size,
            "file_extension": Path(file_path).suffix.lower(),
            "created_time": datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
            "modified_time": datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
            "processing_time": datetime.now().isoformat()
        }

    def _save_json(self, data: Dict, output_path: str) -> None:
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)

    def _create_error_response(self, error_message: str, file_path: str) -> Dict:
        return {
            "content_type": "error",
            "extraction_success": False,
            "error_message": error_message,
            "file_path": file_path,
            "raw_text": "",
            "structured_content": {},
            "metadata": {
                "file_name": os.path.basename(file_path),
                "processing_time": datetime.now().isoformat(),
            },
        }


import re


def clean_key(key):
    """
    Normalize the key by removing parentheses and inner text,
    stripping spaces, lowercasing, and replacing spaces with underscores.
    """
    # Remove parentheses and content inside (e.g., "years_experience (numerical)" -> "years_experience")
    key = re.sub(r'\(.*?\)', '', key)
    key = key.strip().lower().replace(' ', '_')
    return key


def parse_plain_llm_output(text):
    """
    Parses LLM output text where fields are separated by '- ' (dash + space),
    key and value separated by ' is ',
    handles complex values with dashes, commas, and parentheses inside.

    Returns a dictionary of normalized keys and their values.
    """
    res = {}
    # Split by '- ' (note the space after the dash)
    parts = [p.strip() for p in text.strip().split('- ') if p.strip()]

    for part in parts:
        if " is " in part:
            key, val = part.split(" is ", 1)
            key = clean_key(key)
            val = val.strip()
            if val.lower() == "none":
                val = None
            res[key] = val
        else:
            # If part does not contain ' is ', ignore or handle edge cases here if needed
            pass
    return res


def clean_resume_text(text: str) -> str:
    lines = text.splitlines()
    unique_lines = list(dict.fromkeys(lines))
    unique_text = "\n".join(unique_lines).strip()
    return unique_text


def parse_with_openai(text: str) -> str:
    """
    Parse resume text using OpenAI API instead of local LLM
    """
    prompt_header = (
        "You are an information extractor. Extract ONLY from the provided resume text. "
        "If a field is not explicitly present, output None. Do NOT invent or guess values.\n"
        "Output EXACTLY 8 lines, each in the format: '- <field> is <value>'.\n"
        "Fields in order: full_name, email, phone, years_experience, key_skills, previous_companies, education, projects\n"
        "IMPORTANT RULES:\n"
        "- full_name: Extract the person's actual name (first and last name), not job titles or company names\n"
        "- email: Look for email addresses in format user@domain.com\n"
        "- phone: Look for phone numbers (digits, may include +, -, spaces, parentheses)\n"
        "- years_experience: Extract number of years of work experience (just the number)\n"
        "- key_skills: List technical skills, programming languages, tools, frameworks only. Separate with commas.\n"
        "- previous_companies: List actual company names where person worked. Separate with semicolons.\n"
        "- education: Include degree and institution. Separate multiple with semicolons.\n"
        "- projects: List project names or brief descriptions. Separate with semicolons.\n\n"
        "Resume:\n"
    )

    try:
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",  # You can also use "gpt-4" for better quality but higher cost
            messages=[
                {"role": "system",
                 "content": "You are an expert resume parser that extracts information accurately from resume text."},
                {"role": "user", "content": prompt_header + f'"""\n{text}\n"""'}
            ],
            max_tokens=300,
            temperature=0,  # Set to 0 for consistent outputs
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )

        return response.choices[0].message.content.strip()

    except Exception as e:
        print(f"Error calling OpenAI API: {e}")
        return ""


def parse_with_openai_strict(text: str) -> str:
    """Retry prompt with stronger constraints and better formatting using OpenAI API."""
    prompt_header = (
        "Extract ONLY from the resume text below. Never guess or invent. If a field is missing, write None.\n"
        "Return EXACTLY these 8 lines, one per field, format '- <field> is <value>':\n"
        "full_name, email, phone, years_experience, key_skills, previous_companies, education, projects\n"
        "CRITICAL: full_name should be the person's actual name, not job titles or company names.\n"
        "Rules: Use commas for skills, semicolons for companies/education/projects. Keep values concise.\n"
        "Look carefully for contact information and work experience details.\n\n"
        "Resume:\n"
    )

    try:
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system",
                 "content": "You are an expert resume parser. Extract information strictly from the provided text without making assumptions."},
                {"role": "user", "content": prompt_header + f'"""\n{text}\n"""'}
            ],
            max_tokens=300,
            temperature=0,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )

        return response.choices[0].message.content.strip()

    except Exception as e:
        print(f"Error calling OpenAI API (strict): {e}")
        return ""


def fallback_extract_basic_fields(text: str, current: Dict[str, Optional[str]]) -> Dict[str, Optional[str]]:
    """Fill obvious fields from raw text if missing in current dict."""
    updated = dict(current)
    try:
        if not updated.get("email"):
            m = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
            if m:
                updated["email"] = m.group(0)
        if not updated.get("phone"):
            # Find sequences that look like phone numbers, then validate by digit count to avoid capturing dates like 2022-2026
            candidates = re.findall(r"(?:\+\d{1,3}[\s-]?)?(?:\(?\d{2,4}\)?[\s-]?)?\d[\d\s-]{7,}", text)
            best = None
            for c in candidates:
                digits = re.sub(r"\D", "", c)
                if 10 <= len(digits) <= 15:
                    best = c.strip()
                    break
            if best:
                updated["phone"] = best
        if not updated.get("full_name"):
            # Heuristic 1: pick a plausible name from the first 10 lines (increased from 5)
            first_lines = [ln.strip() for ln in text.splitlines()[:10] if ln.strip()]
            name_found = None
            for ln in first_lines:
                # Look for lines that look like names (2-4 words, mostly letters)
                if re.fullmatch(r"[A-Za-z][A-Za-z\-']+(?:\s+[A-Za-z][A-Za-z\-']+){1,3}", ln):
                    # Avoid job titles, company names, and section headers
                    lower_ln = ln.lower()
                    if not any(word in lower_ln for word in
                               ['engineer', 'developer', 'manager', 'director', 'consultant', 'analyst', 'specialist',
                                'coordinator', 'assistant', 'officer', 'executive', 'president', 'ceo', 'cto', 'cfo',
                                'vp', 'head', 'lead', 'senior', 'junior', 'principal', 'staff', 'associate', 'intern',
                                'trainee', 'apprentice', 'student', 'graduate', 'undergraduate', 'phd', 'masters',
                                'bachelor', 'diploma', 'certificate', 'degree', 'university', 'college', 'institute',
                                'school', 'academy', 'corporation', 'company', 'inc', 'ltd', 'llc', 'corp',
                                'enterprise', 'solutions', 'technologies', 'systems', 'services', 'group', 'team',
                                'department', 'division', 'unit', 'section', 'branch', 'office', 'location', 'address',
                                'contact', 'phone', 'email', 'website', 'linkedin', 'github', 'portfolio', 'resume',
                                'cv', 'profile', 'summary', 'objective', 'experience', 'education', 'skills',
                                'projects', 'certifications', 'awards', 'publications', 'patents', 'references']):
                        name_found = ln
                        break
            # Heuristic 2: derive from email local-part if available
            if not name_found and updated.get("email"):
                local = updated["email"].split("@")[0]
                local = re.sub(r"\d+", " ", local)
                parts = re.split(r"[._-]+", local)
                parts = [p for p in parts if p and p.isalpha() and len(p) > 1]
                if parts:
                    name_found = " ".join(w.capitalize() for w in parts[:3])
            if name_found:
                updated["full_name"] = name_found
    except Exception:
        pass
    return updated


def _parse_sections_from_text(text: str) -> Dict[str, List[str]]:
    """Parse sections delimited by lines like '=== SKILLS ===' into a dict."""
    sections: Dict[str, List[str]] = {}
    current: Optional[str] = None
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("===") and stripped.endswith("==="):
            name = stripped.strip("=").strip().lower()
            current = name
            if current not in sections:
                sections[current] = []
            continue
        if current:
            sections[current].append(line)
    return sections


def fallback_extract_from_sections(raw_text: str, current: Dict[str, Optional[str]]) -> Dict[str, Optional[str]]:
    """Use parsed sections and regex to populate missing fields like skills, education, projects, years, companies."""
    updated = dict(current)
    sections = _parse_sections_from_text(raw_text)

    # skills
    if not updated.get("key_skills"):
        skills_lines = sections.get("skills", [])
        if skills_lines:
            joined = " ".join(skills_lines)
            # remove common labels like "Skills", "Languages", "Technologies and tools"
            joined = re.sub(r"\b(Skills?|Languages?|Technologies(?: and tools)?)\b\s*:?,?", " ", joined,
                            flags=re.IGNORECASE)
            # split on commas/semicolons
            skills = [s.strip() for s in re.split(r"[,;]", joined) if s.strip()]
            if skills:
                # dedupe case-insensitive
                seen = set()
                dedup = []
                for s in skills:
                    key = s.lower()
                    if key not in seen:
                        seen.add(key)
                        dedup.append(s)
                updated["key_skills"] = ", ".join(dedup[:50])

    # education - improved extraction
    if not updated.get("education"):
        edu_lines = [ln.strip() for ln in sections.get("education", []) if ln.strip()]
        if edu_lines:
            # Clean up education lines - remove section headers and format better
            cleaned_edu = []
            for line in edu_lines[:10]:
                # Remove common section headers
                line = re.sub(r'^(EDUCATION|Education|ACADEMIC|Academic|QUALIFICATIONS|Qualifications)\s*:?\s*', '',
                              line, flags=re.IGNORECASE)
                if line.strip() and len(line.strip()) > 3:
                    cleaned_edu.append(line.strip())
            if cleaned_edu:
                updated["education"] = "; ".join(cleaned_edu)

    # projects - improved extraction
    if not updated.get("projects"):
        proj_lines = [ln.strip() for ln in sections.get("projects", []) if ln.strip()]
        if proj_lines:
            # Clean up project lines
            cleaned_proj = []
            for line in proj_lines[:12]:
                # Remove common section headers
                line = re.sub(r'^(PROJECTS|Projects|PROJECT|Project)\s*:?\s*', '', line, flags=re.IGNORECASE)
                if line.strip() and len(line.strip()) > 3:
                    cleaned_proj.append(line.strip())
            if cleaned_proj:
                updated["projects"] = "; ".join(cleaned_proj)

    # previous companies from experience - improved extraction
    if not updated.get("previous_companies"):
        exp_lines = sections.get("experience", [])
        companies: List[str] = []
        for ln in exp_lines:
            # Look for company names after "at", "@", or common company indicators
            m = re.search(r"(?:at|@|with|for)\s+([A-Z][A-Za-z0-9&'\.\- ]{2,})", ln, re.IGNORECASE)
            if m:
                companies.append(m.group(1).strip())
            # Common company suffixes
            m2 = re.search(
                r"([A-Z][A-Za-z0-9&'\.\- ]{2,})(?:\s+(?:Inc\.|Inc|LLC|Ltd\.|Ltd|Technologies|Solutions|Labs|Corp\.|Corp|Pvt\.|Pvt|Limited|Company|Co\.|Co))",
                ln)
            if m2:
                companies.append(m2.group(0).strip())
            # Look for standalone company names (capitalized words)
            m3 = re.search(r"([A-Z][A-Za-z0-9&'\.\- ]{3,})(?:\s+[A-Z][A-Za-z0-9&'\.\- ]{2,}){1,3}", ln)
            if m3 and not any(word in ln.lower() for word in ['university', 'college', 'institute', 'school']):
                companies.append(m3.group(0).strip())
        # Deduplicate
        if companies:
            seen = set()
            unique = []
            for c in companies:
                key = c.lower()
                if key not in seen and len(c) > 2:
                    seen.add(key)
                    unique.append(c)
            updated["previous_companies"] = "; ".join(unique[:10])

    # years of experience - improved extraction
    if not updated.get("years_experience"):
        years = []
        # Look for various patterns of years of experience
        patterns = [
            r"(\d+(?:\.\d+)?)\s*(?:\+\s*)?(?:years?|yrs?)\s*(?:in\s+)?(?:the\s+)?(?:field|industry|domain)",
            r"(\d+)\s*(?:\+\s*)?(?:years?|yrs?)"
        ]
        for pattern in patterns:
            for m in re.finditer(pattern, raw_text, flags=re.IGNORECASE):
                try:
                    years.append(float(m.group(1)))
                except Exception:
                    continue
        if years:
            best = max(years)
            updated["years_experience"] = str(int(best)) if best.is_integer() else str(best)

    return updated


def enforce_source_consistency(raw_text: str, fields: Dict[str, Optional[str]]) -> Dict[str, Optional[str]]:
    """Null out values that do not appear in the source text to avoid hallucinations."""
    text_low = raw_text.lower()
    checked = dict(fields)

    def appears(val: Optional[str]) -> bool:
        if not val or val == "None":
            return False
        v = str(val).strip()
        if not v:
            return False
        # accept if any token (>=3 letters) appears in source
        tokens = [t for t in re.split(r"\W+", v.lower()) if len(t) >= 3]
        return any(t in text_low for t in tokens)

    # Only enforce for fields prone to hallucination
    for key in ["full_name", "previous_companies", "education", "projects", "key_skills"]:
        val = checked.get(key)
        if val and not appears(val):
            checked[key] = None

    return checked


def merge_llm_dicts(dict_list: List[Dict[str, Optional[str]]]) -> Dict[str, Optional[str]]:
    """
    Merge parsed chunk dictionaries into a single dictionary.
    - For years_experience: take the max numeric value seen
    - For scalar identity fields (full_name, email, phone): take the first non-empty
    - For text fields (skills, companies, education, projects): merge unique snippets
    """
    merged: Dict[str, Optional[str]] = {
        "full_name": None,
        "email": None,
        "phone": None,
        "years_experience": None,
        "key_skills": None,
        "previous_companies": None,
        "education": None,
        "projects": None,
    }

    def pick_first(current_val: Optional[str], new_val: Optional[str]) -> Optional[str]:
        return current_val if (current_val not in [None, "", "None"]) else (
            new_val if new_val not in [None, "", "None"] else current_val)

    def merge_text(current_val: Optional[str], new_val: Optional[str]) -> Optional[str]:
        parts: List[str] = []
        if current_val and current_val != "None":
            parts.extend([p.strip() for p in str(current_val).split("; ") if p.strip()])
        if new_val and new_val != "None":
            parts.extend([p.strip() for p in str(new_val).split("; ") if p.strip()])
        # dedupe while preserving order
        seen = set()
        unique = []
        for p in parts:
            if p not in seen:
                seen.add(p)
                unique.append(p)
        return "; ".join(unique) if unique else current_val

    def merge_skills(current_val: Optional[str], new_val: Optional[str]) -> Optional[str]:
        skills: List[str] = []
        if current_val and current_val != "None":
            skills.extend([s.strip() for s in str(current_val).split(",") if s.strip()])
        if new_val and new_val != "None":
            skills.extend([s.strip() for s in str(new_val).split(",") if s.strip()])
        # dedupe case-insensitive but keep original casing of first occurrence
        seen_lower = set()
        unique_skills = []
        for s in skills:
            key = s.lower()
            if key not in seen_lower:
                seen_lower.add(key)
                unique_skills.append(s)
        return ", ".join(unique_skills) if unique_skills else current_val

    max_years: Optional[float] = None

    for item in dict_list:
        if not item:
            continue
        merged["full_name"] = pick_first(merged["full_name"], item.get("full_name"))
        merged["email"] = pick_first(merged["email"], item.get("email"))
        merged["phone"] = pick_first(merged["phone"], item.get("phone"))

        y = item.get("years_experience")
        if isinstance(y, str):
            try:
                # extract first number if present
                match = re.search(r"\d+(?:\.\d+)?", y)
                if match:
                    val = float(match.group(0))
                    max_years = val if (max_years is None or val > max_years) else max_years
            except Exception:
                pass
        elif isinstance(y, (int, float)):
            val = float(y)
            max_years = val if (max_years is None or val > max_years) else max_years

        merged["key_skills"] = merge_skills(merged["key_skills"], item.get("key_skills"))
        merged["previous_companies"] = merge_text(merged["previous_companies"], item.get("previous_companies"))
        merged["education"] = merge_text(merged["education"], item.get("education"))
        merged["projects"] = merge_text(merged["projects"], item.get("projects"))

    if max_years is not None:
        merged["years_experience"] = str(int(max_years)) if max_years.is_integer() else str(max_years)

    # Normalize explicit "None" strings to actual None
    for k, v in list(merged.items()):
        if isinstance(v, str) and v.strip().lower() == "none":
            merged[k] = None

    return merged


def chunk_and_parse_with_openai(full_text: str) -> (Dict[str, Optional[str]], str, List[Dict[str, str]]):
    """
    Split resume text into chunks, run OpenAI per chunk, parse each output, and merge results.
    Returns (merged_dict, combined_plain_text_output, debug_info).
    """
    # Auto-bypass chunking when text is short enough (approx <= 3k chars ~ 1k tokens)
    if len(full_text) <= 3000:
        chunks: List[str] = [full_text]
    else:
        chunks: List[str] = text_splitter.split_text(full_text)
    parsed_dicts: List[Dict[str, Optional[str]]] = []
    outputs: List[str] = []
    debug_info: List[Dict[str, str]] = []

    for idx, chunk in enumerate(chunks):
        out = parse_with_openai(chunk)
        try:
            parsed = parse_plain_llm_output(out)
        except Exception:
            parsed = {}
        # Retry if no ' is ' lines parsed
        if not parsed or all(v is None for v in parsed.values()):
            out_retry = parse_with_openai_strict(chunk)
            try:
                parsed_retry = parse_plain_llm_output(out_retry)
            except Exception:
                parsed_retry = {}
            # prefer retry if it yielded anything
            if parsed_retry:
                out = out_retry
                parsed = parsed_retry
        outputs.append(out)
        parsed_dicts.append(parsed)

        prompt_preview = (chunk[:300] + "...") if len(chunk) > 300 else chunk
        print(f"\n[OpenAI  {idx + 1}/{len(chunks)}] Input preview:\n{prompt_preview}")
        print(f"[OpenAI  {idx + 1}] Output:\n{out}")
        print(f"[OpenAI  {idx + 1}] Parsed dict: {parsed}")
        debug_info.append({
            "chunk_index": str(idx),
            "chunk_text": chunk[:1200],
            "output_text": out[:1200],
            "parsed": json.dumps(parsed, ensure_ascii=False)
        })

    merged = merge_llm_dicts(parsed_dicts)
    combined_text_lines = [f"- {k} is {v}" for k, v in merged.items() if v is not None]
    combined_text = "\n".join(combined_text_lines)
    return merged, combined_text, debug_info


def build_and_store_scores(scores_list, db):
    """
    Aggregates scores by job title, sorts them, removes duplicates,
    then upserts the sorted list into the 'score' collection.
    """

    score_collection = db["score"]

    # Group scores by job_title
    grouped = defaultdict(list)
    for entry in scores_list:
        grouped[entry["job_title"]].append({
            "filename": entry["filename"],  # Make sure this key is 'filename'
            "score": entry["score"]
        })

    for job_title, resumes in grouped.items():
        # Remove duplicate resumes per job, by keeping highest score
        seen = {}
        for r in resumes:
            fn = r["filename"]
            if fn not in seen or r["score"] > seen[fn]["score"]:
                seen[fn] = r
        unique_resumes = list(seen.values())

        # Sort resumes descending by score
        sorted_resumes = sorted(unique_resumes, key=lambda x: x["score"], reverse=True)

        # Debug print: confirm data correctness
        print(f"Updating job '{job_title}' with top resumes:")
        for resume in sorted_resumes:
            print(f" - {resume['filename']}: {resume['score']}")

        # Upsert (insert or update) the document in MongoDB
        score_collection.update_one(
            {"job_title": job_title},
            {"$set": {
                "top_resumes": sorted_resumes,
                "updated_at": datetime.now().isoformat()
            }},
            upsert=True
        )


def get_job_description_input():
    jd = input("\nPaste/type your job description or search text, then press Enter:\n")
    return jd.strip()


def fetch_job_descriptions():
    job_collection = db["job_descriptions"]
    # Fetch all job descriptions as list of dicts
    return list(job_collection.find())


def store_match_score(resume_filename, job_title, score, name="", email=""):
    match_collection = db["match_results"]
    # Insert or update score for resume-job pair
    match_collection.update_one(
        {"filename": resume_filename, "job_title": job_title},
        {"$set": {
            "score": score,
            "name": name,
            "email": email
        }},
        upsert=True
    )


def calculate_semantic_similarity(resume_text: str, job_description: str, model) -> float:
    """Calculate semantic similarity using sentence transformers."""
    try:
        resume_embedding = model.encode(resume_text, convert_to_tensor=True)
        job_embedding = model.encode(job_description, convert_to_tensor=True)
        similarity = util.pytorch_cos_sim(resume_embedding, job_embedding).item()
        return round(similarity, 3)
    except Exception as e:
        print(f"Error calculating semantic similarity: {e}")
        return 0.0


def enhanced_parse_with_openai(text: str) -> str:
    """Enhanced OpenAI parsing with better prompts for improved accuracy."""
    prompt_header = (
        "You are an expert resume parser. Extract ONLY from the provided resume text. "
        "If a field is not explicitly present, output None. Do NOT invent or guess values.\n"
        "Output EXACTLY 8 lines, each in the format: '- <field> is <value>'.\n"
        "Fields in order: full_name, email, phone, years_experience, key_skills, previous_companies, education, projects\n"
        "CRITICAL RULES:\n"
        "- full_name: Extract the person's actual name (first and last name), NOT job titles or company names\n"
        "- email: Look for email addresses in format user@domain.com\n"
        "- phone: Look for phone numbers (digits, may include +, -, spaces, parentheses)\n"
        "- years_experience: Extract number of years of work experience (just the number)\n"
        "- key_skills: List technical skills, programming languages, tools, frameworks only. Separate with commas.\n"
        "- previous_companies: List actual company names where person worked. Separate with semicolons.\n"
        "- education: Include degree and institution. Separate multiple with semicolons.\n"
        "- projects: List project names or brief descriptions. Separate with semicolons.\n\n"
        "Resume:\n"
    )

    try:
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system",
                 "content": "You are an expert resume parser that extracts information accurately from resume text."},
                {"role": "user", "content": prompt_header + f'"""\n{text}\n"""'}
            ],
            max_tokens=350,
            temperature=0,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )

        return response.choices[0].message.content.strip()

    except Exception as e:
        print(f"Error calling OpenAI API (enhanced): {e}")
        return ""


def enhanced_fallback_extract_basic_fields(text: str, current: Dict[str, Optional[str]]) -> Dict[str, Optional[str]]:
    """Enhanced fallback extraction with better name detection."""
    updated = dict(current)
    try:
        if not updated.get("email"):
            m = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
            if m:
                updated["email"] = m.group(0)
        if not updated.get("phone"):
            candidates = re.findall(r"(?:\+\d{1,3}[\s-]?)?(?:\(?\d{2,4}\)?[\s-]?)?\d[\d\s-]{7,}", text)
            best = None
            for c in candidates:
                digits = re.sub(r"\D", "", c)
                if 10 <= len(digits) <= 15:
                    best = c.strip()
                    break
            if best:
                updated["phone"] = best
        if not updated.get("full_name"):
            # Enhanced name detection from first 15 lines
            first_lines = [ln.strip() for ln in text.splitlines()[:15] if ln.strip()]
            name_found = None
            for ln in first_lines:
                if re.fullmatch(r"[A-Za-z][A-Za-z\-']+(?:\s+[A-Za-z][A-Za-z\-']+){1,3}", ln):
                    lower_ln = ln.lower()
                    # Enhanced filtering for job titles and company names
                    if not any(word in lower_ln for word in
                               ['engineer', 'developer', 'manager', 'director', 'consultant', 'analyst', 'specialist',
                                'coordinator', 'assistant', 'officer', 'executive', 'president', 'ceo', 'cto', 'cfo',
                                'vp', 'head', 'lead', 'senior', 'junior', 'principal', 'staff', 'associate', 'intern']):
                        name_found = ln
                        break
            # Enhanced email-based name extraction
            if not name_found and updated.get("email"):
                local = updated["email"].split("@")[0]
                local = re.sub(r"\d+", " ", local)
                parts = re.split(r"[._-]+", local)
                parts = [p for p in parts if p and p.isalpha() and len(p) > 1]
                if parts:
                    name_found = " ".join(w.capitalize() for w in parts[:3])
            if name_found:
                updated["full_name"] = name_found
    except Exception:
        pass
    return updated


def enhanced_fallback_extract_from_sections(raw_text: str, current: Dict[str, Optional[str]]) -> Dict[
    str, Optional[str]]:
    """Enhanced section-based extraction with better cleaning."""
    updated = dict(current)
    sections = _parse_sections_from_text(raw_text)

    # Enhanced skills extraction
    if not updated.get("key_skills"):
        skills_lines = sections.get("skills", [])
        if skills_lines:
            joined = " ".join(skills_lines)
            joined = re.sub(r"\b(Skills?|Languages?|Technologies(?: and tools)?)\b\s*:?,?", " ", joined,
                            flags=re.IGNORECASE)
            skills = [s.strip() for s in re.split(r"[,;]", joined) if s.strip()]
            if skills:
                seen = set()
                dedup = []
                for s in skills:
                    key = s.lower()
                    if key not in seen:
                        seen.add(key)
                        dedup.append(s)
                updated["key_skills"] = ", ".join(dedup[:50])

    # Enhanced education extraction
    if not updated.get("education"):
        edu_lines = [ln.strip() for ln in sections.get("education", []) if ln.strip()]
        if edu_lines:
            cleaned_edu = []
            for line in edu_lines[:10]:
                line = re.sub(
                    r'^(EDUCATION|Education|ACADEMIC|Academic|QUALIFICATIONS|Qualifications|CERTIFICATES|Certificates)\s*:?\s*',
                    '', line, flags=re.IGNORECASE)
                if line.strip() and len(line.strip()) > 3:
                    cleaned_edu.append(line.strip())
            if cleaned_edu:
                updated["education"] = "; ".join(cleaned_edu)

    # Enhanced projects extraction
    if not updated.get("projects"):
        proj_lines = [ln.strip() for ln in sections.get("projects", []) if ln.strip()]
        if proj_lines:
            cleaned_proj = []
            for line in proj_lines[:12]:
                line = re.sub(r'^(PROJECTS|Projects|PROJECT|Project)\s*:?\s*', '', line, flags=re.IGNORECASE)
                if line.strip() and len(line.strip()) > 3:
                    cleaned_proj.append(line.strip())
            if cleaned_proj:
                updated["projects"] = "; ".join(cleaned_proj)

    # Enhanced company extraction
    if not updated.get("previous_companies"):
        exp_lines = sections.get("experience", [])
        companies: List[str] = []
        for ln in exp_lines:
            m = re.search(r"(?:at|@|with|for)\s+([A-Z][A-Za-z0-9&'\.\- ]{2,})", ln, re.IGNORECASE)
            if m:
                companies.append(m.group(1).strip())
            m2 = re.search(
                r"([A-Z][A-Za-z0-9&'\.\- ]{2,})(?:\s+(?:Inc\.|Inc|LLC|Ltd\.|Ltd|Technologies|Solutions|Labs|Corp\.|Corp|Pvt\.|Pvt|Limited|Company|Co\.|Co))",
                ln)
            if m2:
                companies.append(m2.group(0).strip())
            m3 = re.search(r"([A-Z][A-Za-z0-9&'\.\- ]{3,})(?:\s+[A-Z][A-Za-z0-9&'\.\- ]{2,}){1,3}", ln)
            if m3 and not any(word in ln.lower() for word in ['university', 'college', 'institute', 'school']):
                companies.append(m3.group(0).strip())
        if companies:
            seen = set()
            unique = []
            for c in companies:
                key = c.lower()
                if key not in seen and len(c) > 2:
                    seen.add(key)
                    unique.append(c)
            updated["previous_companies"] = "; ".join(unique[:10])

    # Enhanced years extraction
    if not updated.get("years_experience"):
        years = []
        patterns = [
            r"(\d+(?:\.\d+)?)\s*(?:\+\s*)?(?:years?|yrs?)\s*(?:of\s+)?(?:experience|exp)",
            r"(?:experience|exp)\s*(?:of\s+)?(\d+(?:\.\d+)?)\s*(?:\+\s*)?(?:years?|yrs?)",
            r"(\d+(?:\.\d+)?)\s*(?:\+\s*)?(?:years?|yrs?)\s*(?:in\s+)?(?:the\s+)?(?:field|industry|domain)",
            r"(\d+)\s*(?:\+\s*)?(?:years?|yrs?)"
        ]
        for pattern in patterns:
            for m in re.finditer(pattern, raw_text, flags=re.IGNORECASE):
                try:
                    years.append(float(m.group(1)))
                except Exception:
                    continue
        if years:
            best = max(years)
            updated["years_experience"] = str(int(best)) if best.is_integer() else str(best)

    return updated


def enhanced_parse_with_openai_strict(text: str) -> str:
    """Enhanced strict OpenAI parsing with stronger constraints."""
    prompt_header = (
        "Extract ONLY from the resume text below. Never guess or invent. If a field is missing, write None.\n"
        "Return EXACTLY these 8 lines, one per field, format '- <field> is <value>':\n"
        "full_name, email, phone, years_experience, key_skills, previous_companies, education, projects\n"
        "CRITICAL: full_name should be the person's actual name, not job titles or company names.\n"
        "Rules: Use commas for skills, semicolons for companies/education/projects. Keep values concise.\n"
        "Look carefully for contact information and work experience details.\n\n"
        "Resume:\n"
    )

    try:
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system",
                 "content": "You are an expert resume parser. Extract information strictly from the provided text without making assumptions."},
                {"role": "user", "content": prompt_header + f'"""\n{text}\n"""'}
            ],
            max_tokens=350,
            temperature=0,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )

        return response.choices[0].message.content.strip()

    except Exception as e:
        print(f"Error calling OpenAI API (enhanced strict): {e}")
        return ""


def enhanced_chunk_and_parse_with_openai(full_text: str) -> (Dict[str, Optional[str]], str, List[Dict[str, str]]):
    """
    Enhanced chunking and parsing with better OpenAI prompts.
    Returns (merged_dict, combined_plain_text_output, debug_info).
    """
    # Auto-bypass chunking when text is short enough (approx <= 3k chars ~ 1k tokens)
    if len(full_text) <= 3000:
        chunks: List[str] = [full_text]
    else:
        chunks: List[str] = text_splitter.split_text(full_text)
    parsed_dicts: List[Dict[str, Optional[str]]] = []
    outputs: List[str] = []
    debug_info: List[Dict[str, str]] = []

    for idx, chunk in enumerate(chunks):
        out = enhanced_parse_with_openai(chunk)
        try:
            parsed = parse_plain_llm_output(out)
        except Exception:
            parsed = {}
        # Retry if no ' is ' lines parsed
        if not parsed or all(v is None for v in parsed.values()):
            out_retry = enhanced_parse_with_openai_strict(chunk)
            try:
                parsed_retry = parse_plain_llm_output(out_retry)
            except Exception:
                parsed_retry = {}
            # prefer retry if it yielded anything
            if parsed_retry:
                out = out_retry
                parsed = parsed_retry
        outputs.append(out)
        parsed_dicts.append(parsed)

        prompt_preview = (chunk[:300] + "...") if len(chunk) > 300 else chunk
        print(f"\n[ENHANCED OpenAI CHUNK {idx + 1}/{len(chunks)}] Input preview:\n{prompt_preview}")
        print(f"[ENHANCED OpenAI CHUNK {idx + 1}] Output:\n{out}")
        print(f"[ENHANCED OpenAI CHUNK {idx + 1}] Parsed dict: {parsed}")
        debug_info.append({
            "chunk_index": str(idx),
            "chunk_text": chunk[:1200],
            "output_text": out[:1200],
            "parsed": json.dumps(parsed, ensure_ascii=False)
        })

    merged = merge_llm_dicts(parsed_dicts)
    combined_text_lines = [f"- {k} is {v}" for k, v in merged.items() if v is not None]
    combined_text = "\n".join(combined_text_lines)
    return merged, combined_text, debug_info


def use_enhanced_parsing_with_openai(cleaned_text: str, raw_text: str) -> (
Dict[str, Optional[str]], str, List[Dict[str, str]]):
    """
    Use enhanced parsing functions with OpenAI for better accuracy.
    This function can be called instead of the original chunk_and_parse_with_llm.
    """
    # Use enhanced OpenAI parsing
    merged_fields, llm_output_str, llm_debug = enhanced_chunk_and_parse_with_openai(cleaned_text)

    # Use enhanced fallback extraction
    merged_fields = enhanced_fallback_extract_basic_fields(cleaned_text, merged_fields)
    merged_fields = enhanced_fallback_extract_from_sections(raw_text, merged_fields)
    merged_fields = enforce_source_consistency(raw_text, merged_fields)

    return merged_fields, llm_output_str, llm_debug


def main():
    input_files = [
        f for f in os.listdir("data")
        if f.endswith(('.pdf', '.docx', '.doc'))
    ]

    if not input_files:
        print("📂 No resumes found in /data folder.")
        return

    converter = DocumentToJSONConverter()

    # No need to initialize local pipeline anymore - using OpenAI API
    print("🌐 Using OpenAI API for resume parsing...")

    # Use GPU if available for sentence transformer, else CPU
    scoring_model = SentenceTransformer('all-MiniLM-L6-v2')

    job_descriptions = fetch_job_descriptions()
    if not job_descriptions:
        print("❌ No job descriptions found in the database.")
        return

    unique_scores = {}

    for filename in input_files:
        path = os.path.join("data", filename)
        print(f"\n📄 Processing: {filename}")

        result = converter.convert_to_json(path)
        if not result.get("extraction_success"):
            print("❌ Extraction failed:", result.get("error_message", "Unknown error"))
            continue

        cleaned = clean_resume_text(result["raw_text"])

        # 1. Get raw OpenAI output using chunking and merge results
        merged_fields, llm_output_str, llm_debug = chunk_and_parse_with_openai(cleaned)
        print(f"Raw OpenAI output for file '{filename}':")
        print(llm_output_str)

        # 1b. Fallback: fill obvious fields from raw text if OpenAI missed them
        merged_fields = fallback_extract_basic_fields(cleaned, merged_fields)
        # 1c. Fallback from structured sections/regex for remaining fields
        merged_fields = fallback_extract_from_sections(result["raw_text"], merged_fields)
        # 1d. Enforce values exist in source to avoid copied examples
        merged_fields = enforce_source_consistency(result["raw_text"], merged_fields)

        # 2. Use merged_fields directly (already parsed)
        llm_data = merged_fields if isinstance(merged_fields, dict) else {}

        # 3. Debug prints: show parsed data and keys
        print(f"Parsed OpenAI data for '{filename}': {llm_data}")
        print(f"Parsed keys in OpenAI output for '{filename}': {list(llm_data.keys())}")

        # 4. Insert resume metadata + raw output if not exists
        if not collection.find_one({"filename": filename}):
            collection.insert_one({
                "filename": filename,
                "filepath": path,
                "raw_text": result["raw_text"],
                "llm_output": llm_output_str,
                "llm_debug": llm_debug[:10],
                "extraction_method": result.get("extraction_method"),
                "metadata": result.get("metadata", {}),
                "created_at": datetime.now().isoformat()
            })
            print("📥 Saved to MongoDB Atlas ✅")
        else:
            print(f"⚠️ {filename} already exists in MongoDB. Skipping upload.")

        # 5. Prepare document for graphs_llm collection (use get with default None)
        clean_doc = {
            "resume_id": filename,
            "full_name": llm_data.get("full_name"),
            "email": llm_data.get("email"),
            "phone": llm_data.get("phone"),
            "years_experience": llm_data.get("years_experience"),
            "key_skills": llm_data.get("key_skills"),
            "previous_companies": llm_data.get("previous_companies"),
            "education": llm_data.get("education"),
            "projects": llm_data.get("projects"),
            "filename": filename,
            "created_at": datetime.now().isoformat(),
            "job_matches": [],
            "llm_debug": llm_debug[:10],
        }
        for job in job_descriptions:
            job_title = job.get("title", "Unknown Job")
            job_desc_text = job.get("description", "")

            # Use simple semantic similarity scoring
            resume_text = f"{llm_data.get('key_skills', '')} {llm_data.get('projects', '')} {llm_data.get('previous_companies', '')}"
            score = calculate_semantic_similarity(resume_text, job_desc_text, scoring_model)

            clean_doc["job_matches"].append({
                "job_title": job_title,
                "job_description": job_desc_text,
                "score": score
            })

            # Store match score
            store_match_score(filename, job_title, score,
                              name=llm_data.get("full_name", ""),
                              email=llm_data.get("email", ""))

        print(f"Storing into graphs_llm collection: {clean_doc}")

        # 6. Upsert the parsed resume info
        graphs_collection.update_one(
            {"filename": filename},
            {"$set": clean_doc},
            upsert=True
        )
        print("Debug: Fields just before MongoDB upsert:")
        for key in ["years_experience", "projects"]:
            print(f"{key}: {clean_doc.get(key)}")
        print("Full document:", clean_doc)
        # Step 1: Immediately check what was actually stored in MongoDB
        stored_doc = graphs_collection.find_one({"filename": filename})
        print("🔍 Stored in MongoDB (graphs_llm) just after upsert:", stored_doc)

        print(f"🗂 Parsed OpenAI data stored in '{graphs_collection}' collection.")

        # 7. Score resume against each job description
        for job in job_descriptions:
            job_title = job.get("title", "Unknown Job")
            job_desc_text = job.get("description", "")

            # Use simple semantic similarity scoring
            resume_text = f"{llm_data.get('key_skills', '')} {llm_data.get('projects', '')} {llm_data.get('previous_companies', '')}"
            score = calculate_semantic_similarity(resume_text, job_desc_text, scoring_model)
            print(f"Score for resume '{filename}' against job '{job_title}': {score}")

            # Store or update match score in 'match_results' collection
            store_match_score(filename, job_title, score)

            # Keep track uniquely for final reporting
            unique_scores[(filename, job_title)] = {
                "score": score,
                "filename": filename,
                "job_title": job_title
            }

    # 8. After processing all resumes, prepare scores list
    scores = list(unique_scores.values())

    # 9. Print top matches per job description
    print("\n🏆 Top Matches per Job Description:")
    for job in job_descriptions:
        job_title = job.get("title", "Unknown Job")
        filtered_scores = [s for s in scores if s["job_title"] == job_title]
        ranked = sorted(filtered_scores, key=lambda x: x["score"], reverse=True)
        print(f"\nJob: {job_title}")
        for i, res in enumerate(ranked[:5], 1):
            print(f"    {i}. {res['filename']} — Score: {res['score']}")

    # 10. Final step: build and store all scores
    build_and_store_scores(scores, db)


if __name__ == "__main__":
    main()
